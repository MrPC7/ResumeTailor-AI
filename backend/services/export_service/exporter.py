from __future__ import annotations

import re
from dataclasses import dataclass
from io import BytesIO

import fitz
from docx import Document

from schemas.extract_resume import ExtractResumeResponse
from schemas.export import ExportFormat


class ResumeExportError(Exception):
    """Raised when export generation fails."""


@dataclass(frozen=True)
class ExportFile:
    content: bytes
    media_type: str
    extension: str


class ResumeExporter:
    def export(self, resume: ExtractResumeResponse, export_format: ExportFormat) -> ExportFile:
        if export_format == ExportFormat.PDF:
            return ExportFile(
                content=self._build_pdf(resume),
                media_type="application/pdf",
                extension="pdf",
            )

        if export_format == ExportFormat.DOCX:
            return ExportFile(
                content=self._build_docx(resume),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                extension="docx",
            )

        raise ResumeExportError("Unsupported export format.")

    @staticmethod
    def build_file_base_name(resume_name: str | None, requested_name: str | None) -> str:
        raw_name = (requested_name or resume_name or "resume").strip()
        sanitized = re.sub(r"[^A-Za-z0-9_-]", "_", raw_name)
        normalized = re.sub(r"_+", "_", sanitized).strip("_")
        return normalized or "resume"

    @staticmethod
    def _build_pdf(resume: ExtractResumeResponse) -> bytes:
        MARGIN = 48.0
        NAME_SIZE = 18.0
        HEADING_SIZE = 14.0
        BODY_SIZE = 11.0
        NAME_LINE_HEIGHT = 24.0
        HEADING_LINE_HEIGHT = 20.0
        BODY_LINE_HEIGHT = 15.0
        SECTION_GAP = 12.0
        FONT_BODY = "helv"
        FONT_BOLD = "hebo"

        document = fitz.open()
        try:
            page = document.new_page()
            y = MARGIN
            left = MARGIN
            right_limit = page.rect.width - MARGIN
            bottom_limit = page.rect.height - MARGIN

            def _ensure_space(needed: float) -> None:
                nonlocal page, y
                if y + needed > bottom_limit:
                    page = document.new_page()
                    y = MARGIN

            def _wrap_text(text: str, fontsize: float, fontname: str) -> list[str]:
                """Word-wrap text to fit within page margins."""
                max_width = right_limit - left
                words = text.split()
                lines: list[str] = []
                current_line = ""
                for word in words:
                    test = f"{current_line} {word}".strip()
                    tw = fitz.get_text_length(test, fontname=fontname, fontsize=fontsize)
                    if tw <= max_width:
                        current_line = test
                    else:
                        if current_line:
                            lines.append(current_line)
                        current_line = word
                if current_line:
                    lines.append(current_line)
                return lines or [" "]

            def _write(text: str, fontsize: float, fontname: str, line_height: float) -> None:
                nonlocal y
                for raw_line in text.split("\n"):
                    wrapped = _wrap_text(raw_line or " ", fontsize, fontname)
                    for wl in wrapped:
                        _ensure_space(line_height)
                        page.insert_text((left, y), wl, fontsize=fontsize, fontname=fontname)
                        y += line_height

            # Name
            _write(resume.name or "Customized Resume", NAME_SIZE, FONT_BOLD, NAME_LINE_HEIGHT)

            # Contact
            contact_parts = [p for p in [resume.email, resume.phone] if p]
            if contact_parts:
                _write(" | ".join(contact_parts), BODY_SIZE, FONT_BODY, BODY_LINE_HEIGHT)

            # Summary
            if resume.summary:
                y += SECTION_GAP
                _write("Summary", HEADING_SIZE, FONT_BOLD, HEADING_LINE_HEIGHT)
                _write(resume.summary, BODY_SIZE, FONT_BODY, BODY_LINE_HEIGHT)

            # Skills
            if resume.skills:
                y += SECTION_GAP
                _write("Skills", HEADING_SIZE, FONT_BOLD, HEADING_LINE_HEIGHT)
                _write(", ".join(s.strip() for s in resume.skills if s.strip()), BODY_SIZE, FONT_BODY, BODY_LINE_HEIGHT)

            # Experience
            if resume.experience:
                y += SECTION_GAP
                _write("Experience", HEADING_SIZE, FONT_BOLD, HEADING_LINE_HEIGHT)
                for item in resume.experience:
                    header = " - ".join(p for p in [item.position, item.company] if p)
                    if header:
                        _ensure_space(BODY_LINE_HEIGHT)
                        _write(header, BODY_SIZE, FONT_BOLD, BODY_LINE_HEIGHT)
                    if item.duration:
                        _write(item.duration, BODY_SIZE, FONT_BODY, BODY_LINE_HEIGHT)
                    if item.description:
                        _write(item.description, BODY_SIZE, FONT_BODY, BODY_LINE_HEIGHT)
                    y += 4.0  # spacing between entries

            # Education
            if resume.education:
                y += SECTION_GAP
                _write("Education", HEADING_SIZE, FONT_BOLD, HEADING_LINE_HEIGHT)
                for item in resume.education:
                    edu_line = ", ".join(p for p in [item.degree, item.institution, item.year] if p)
                    if edu_line:
                        _write(edu_line, BODY_SIZE, FONT_BODY, BODY_LINE_HEIGHT)

            # Projects
            if resume.projects:
                y += SECTION_GAP
                _write("Projects", HEADING_SIZE, FONT_BOLD, HEADING_LINE_HEIGHT)
                for item in resume.projects:
                    if item.name:
                        _write(item.name, BODY_SIZE, FONT_BOLD, BODY_LINE_HEIGHT)
                    if item.description:
                        _write(item.description, BODY_SIZE, FONT_BODY, BODY_LINE_HEIGHT)
                    if item.technologies:
                        tech_line = ", ".join(t.strip() for t in item.technologies if t.strip())
                        if tech_line:
                            _write(f"Technologies: {tech_line}", BODY_SIZE, FONT_BODY, BODY_LINE_HEIGHT)
                    y += 4.0

            return document.tobytes(garbage=4, deflate=True)
        except ResumeExportError:
            raise
        except Exception as exc:
            raise ResumeExportError("Unable to generate PDF file.") from exc
        finally:
            document.close()

    @staticmethod
    def _build_docx(resume: ExtractResumeResponse) -> bytes:
        try:
            document = Document()
            document.add_heading(resume.name or "Customized Resume", level=0)

            if resume.email or resume.phone:
                contact_parts = [part for part in [resume.email, resume.phone] if part]
                document.add_paragraph(" | ".join(contact_parts))

            if resume.summary:
                document.add_heading("Summary", level=1)
                document.add_paragraph(resume.summary)

            if resume.skills:
                document.add_heading("Skills", level=1)
                document.add_paragraph(", ".join(skill.strip() for skill in resume.skills if skill.strip()))

            if resume.experience:
                document.add_heading("Experience", level=1)
                for item in resume.experience:
                    header = " - ".join(part for part in [item.position, item.company] if part)
                    if header:
                        document.add_paragraph(header)
                    if item.duration:
                        document.add_paragraph(item.duration)
                    if item.description:
                        document.add_paragraph(item.description)

            if resume.education:
                document.add_heading("Education", level=1)
                for item in resume.education:
                    edu_line = ", ".join(
                        part for part in [item.degree, item.institution, item.year] if part
                    )
                    if edu_line:
                        document.add_paragraph(edu_line)

            if resume.projects:
                document.add_heading("Projects", level=1)
                for item in resume.projects:
                    if item.name:
                        document.add_paragraph(item.name)
                    if item.description:
                        document.add_paragraph(item.description)
                    if item.technologies:
                        tech_line = ", ".join(
                            tech.strip() for tech in item.technologies if tech.strip()
                        )
                        if tech_line:
                            document.add_paragraph(f"Technologies: {tech_line}")

            stream = BytesIO()
            document.save(stream)
            return stream.getvalue()
        except Exception as exc:
            raise ResumeExportError("Unable to generate DOCX file.") from exc


resume_exporter = ResumeExporter()
